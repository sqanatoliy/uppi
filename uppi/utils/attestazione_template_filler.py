import shutil
from docx import Document
from docx.enum.text import WD_UNDERLINE
from pathlib import Path

SPACE_TO_UNDERSCORE_RATIO = 1.8

# ----------------------------------------------
# Допоміжна функція:
# ----------------------------------------------
def fill_underscored(text: str | None, length: int) -> str:

    # 1. Якщо даних немає — повертаємо оригінальні підкреслення
    if not text:
        return "_" * length

    text = str(text).strip()

    # 2. Розрахунок центрованих пробілів
    
    # Кількість символів, яку треба заповнити в оригінальній метриці '_'
    original_padding_length = length - len(text)
    
    if original_padding_length <= 0:
        return text

    # Розрахунок загальної кількості пробілів, необхідних для збереження ширини
    # Обов'язково округлюємо (або відкидаємо дробову частину) до цілого числа (int)
    target_space_padding = int(original_padding_length * SPACE_TO_UNDERSCORE_RATIO)

    # Центрування в новій метриці пробілів
    left_padding = target_space_padding // 2
    right_padding = target_space_padding - left_padding
    
    # Тепер множимо рядок на цілі числа, що усуває TypeError
    return " " * left_padding + text + " " * right_padding

def replace_in_paragraph(paragraph, params, underscored):
    """Замінює ключі в параграфі документа."""
    for run in paragraph.runs:
        for key, value in params.items():
            if key in run.text:
                if key == "{{CONDUTTORE_CF}}" and len(str(value)) > 0:
                    # Особливий випадок для CONDUTTORE codice fiscale
                    replacement_text = fill_underscored(value, 0)
                    run.text = run.text.replace(key, replacement_text)
                    run.font.underline = WD_UNDERLINE.SINGLE
                if key in underscored:
                    replacement_text = fill_underscored(value, underscored[key])
                    run.text = run.text.replace(key, replacement_text)
                    run.font.underline = WD_UNDERLINE.SINGLE
                else:
                    replacement = str(value or "")
                    run.text = run.text.replace(key, replacement)

def replace_in_cell(cell, params, underscored):
    """Замінює ключі в комірці таблиці."""
    for paragraph in cell.paragraphs:
        replace_in_paragraph(paragraph, params, underscored)


# ---------------------------------------------------
# Основна функція заповнення шаблону
# params: словник { "{{KEY}}": value }
# underscored: словник { "{{KEY}}": fixed_length }
# ---------------------------------------------------
def fill_attestazione_template(
    template_path: str,
    output_folder: str,
    filename: str,
    params: dict,
    underscored: dict
):
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    out_path = output_folder / filename

    # копія шаблону
    shutil.copy(template_path, out_path)

    # завантажуємо документ
    doc = Document(out_path)

    # обробка всіх параграфів
# параграфи
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, params, underscored)

    # таблиці
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, params, underscored)

    doc.save(out_path)
    return str(out_path)

# ----------------------------------------------
# Фіксована довжина підкреслень для відповідних ключів
underscored = {
    "{{LOCATORE_NOME}}": 40,
    "{{LOCATORE_CF}}": 25,
    "{{LOCATORE_COMUNE_RES}}": 27,
    "{{LOCATORE_VIA}}": 27,
    "{{LOCATORE_CIVICO}}": 4,
    "{{IMMOBILE_COMUNE}}": 24,
    "{{IMMOBILE_VIA}}": 27,
    "{{IMMOBILE_CIVICO}}": 4,
    "{{IMMOBILE_PIANO}}": 4,
    "{{IMMOBILE_INTERNO}}": 4,
    "{{CONTRATTO_DATA}}": 13,
    "{{CONDUTTORE_NOME}}": 27,
    "{{CONDUTTORE_CF}}": 21,
    "{{CONDUTTORE_COMUNE}}": 24,
    "{{CONDUTTORE_VIA}}": 27,
    "{{DECORRENZA_DATA}}": 18,
    "{{REGISTRAZIONE_DATA}}": 13,
    "{{REGISTRAZIONE_NUM}}": 4,
    "{{AGENZIA_ENTRATE_SEDE}}": 25,
}


if __name__ == "__main__":
    BASE_DIR = Path(__file__).resolve().parent.parent.parent
    template_path = BASE_DIR/"attestazione_template/template_attestazione_pescara.docx"

    output_dir = "output"
    filename = "attestazione_filled.docx"



    params = {
        # Dati del locatore
        "{{LOCATORE_NOME}}": "Mario Rossi",
        "{{LOCATORE_CF}}": "RSSMRA80A01H501X",
        "{{LOCATORE_COMUNE_RES}}": "Pescara",
        "{{LOCATORE_VIA}}": "Predazzo",
        "{{LOCATORE_CIVICO}}": 43,

        # Dati dell'immobile
        "{{IMMOBILE_COMUNE}}": "Montesilvano",
        "{{IMMOBILE_VIA}}": "C-so Umberto I",
        "{{IMMOBILE_CIVICO}}": "316",
        "{{IMMOBILE_PIANO}}": "4",
        "{{IMMOBILE_INTERNO}}": "",

        # Dati del contratto
        "{{CONTRATTO_DATA}}": "",

        # Dati del conduttore
        "{{CONDUTTORE_NOME}}": "Biaocchi Giovana",
        "{{CONDUTTORE_CF}}": "BCCGNN44M45G488W",
        "{{CONDUTTORE_COMUNE}}": "",
        "{{CONDUTTORE_VIA}}": "",

        # Dati della registrazione
        "{{DECORRENZA_DATA}}": "15/10/2025",
        "{{REGISTRAZIONE_DATA}}": "",
        "{{REGISTRAZIONE_NUM}}": "",
        # Agenzia delle Entrate sede
        "{{AGENZIA_ENTRATE_SEDE}}": "",

        # elementi immobiliari tipo A e B
        "{{a1}}": "X",
        "{{a2}}": "X",
        "{{b1}}": "X",
        "{{b2}}": "X",
        "{{b3}}": "X",
        "{{b4}}": "X",
        "{{b5}}": "X",
        # elementi immobiliari tipo C
        "{{c1}}": "",
        "{{c2}}": "",
        "{{c3}}": "X",
        "{{c4}}": "X",
        "{{c5}}": "X",
        "{{c6}}": "X",
        "{{c7}}": "X",
        # elementi immobiliari tipo D
        "{{d1}}": "X",
        "{{d2}}": "",
        "{{d3}}": "",
        "{{d4}}": "X",
        "{{d5}}": "X",
        "{{d6}}": "",
        "{{d7}}": "",
        "{{d8}}": "",
        "{{d9}}": "",
        "{{d10}}": "",
        "{{d11}}": "",
        "{{d12}}": "",
        "{{d13}}": "",
        # calcolo numero elementi
        "{{A_CNT}}": "",
        "{{B_CNT}}": "",
        "{{C_CNT}}": "",
        "{{D_CNT}}": "",
        # dati catastali
        # Appartamento
        "{{APP_FOGL}}": "7",
        "{{APP_PART}}": "886",
        "{{APP_SUB}}": "101",
        "{{APP_REND}}": "216,91",
        "{{APP_SCAT}}": "55",
        "{{APP_SRIP}}": "",
        "{{APP_CAT}}": "A/2",
        # Garage/cantina
        "{{GAR_FOGL}}": "",
        "{{GAR_PART}}": "",
        "{{GAR_SUB}}": "",
        "{{GAR_REND}}": "",
        "{{GAR_SCAT}}": "",
        "{{GAR_SRIP}}": "",
        "{{GAR_CAT}}": "",
        # Posto maccina
        "{{PST_FOGL}}": "",
        "{{PST_PART}}": "",
        "{{PST_SUB}}": "",
        "{{PST_REND}}": "",
        "{{PST_SCAT}}": "",
        "{{PST_SRIP}}": "",
        "{{PST_CAT}}": "",
        # Totale superficile
        "{{TOT_SCAT}}": "",
        "{{TOT_SRIP}}": "",
        "{{TOT_CAT}}": "",

        # Calcolo del canone
        #ZONA OMOGENEA
        "{{CAN_ZONA}}": "1",
        # SUB-FASCIA DI APPARTENENZA
        "{{CAN_SUBFASCIA}}": "3",
        # METRI QUADRI
        "{{CAN_MQ}}": "55",
        # CANONE ANNUO AL METRO QUADRO
        "{{CAN_MQ_ANNUO}}": "97,56",
        # CANONE ANNUO TOTALE IMMOBILE LOCATO
        "{{CAN_TOTALE_ANNUO}}": "5365,90",
        # MAGGIORAZIONE IMMOBILE ARREDATO MAX  15%
        "{{CAN_ARREDATO}}": "804,88",
        # MAGGIORAZIONE CLASSE A 8%
        "{{CAN_CLASSE_A}}": "",
        # MAGGIORAZIONE CLASSE B 4%
        "{{CAN_CLASSE_B}}": "",
        # DIMINUZIONE CLASSI ENERGETICHE E (-2%) F (-4%) G (-6%)
        "{{CAN_ENERGY}}": "321,95",
        # MAGGIORAZIONE DURATA CONTRATTO + 5% (quattro anni) – 6% (cinque anni) – 7% (sei e più anni)
        "{{CAN_DURATA}}": "",
        # MAGGIORAZIONE TRANSITORIO 15%
        "{{CAN_TRANSITORIO}}": "",
        # MAGGIORAZIONE STUDENTI UNIVERSITARI 20%
        "{{CAN_STUDENTI}}": "",
        # TOTALE CANONE ANNUO CON VARIAZIONI PERCENTUALI Minimo
        "{{CAN_ANNUO_VAR_MIN}}": "5095,75",
        # TOTALE CANONE ANNUO CON VARIAZIONI PERCENTUALI Massimo
        "{{CAN_ANNUO_VAR_MAX}}": "5848,82",
        # TOTALE CANONE MENSILE CON VARIAZIONI PERCENTUUALI Minimo
        "{{CAN_MENSILE_VAR_MIN}}": "424,64",
        # TOTALE CANONE MENSILE CON VARIAZIONI PERCENTUUALI Massimo
        "{{CAN_MENSILE_VAR_MAX}}": "487,40",
        # TOTALE MENSILE CONCORDATO TRA LE PARTI
        "{{CAN_MENSILE}}": "485,00",
    }

    print(BASE_DIR)
    filled_path = fill_attestazione_template(
        template_path,
        output_dir,
        filename,
        params,
        underscored
    )

    print(f"Filled document saved to: {filled_path}")